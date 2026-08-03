from copy import deepcopy

from docx.shared import RGBColor
from docx.table import _Row


def replace_in_paragraph(paragraph, old, new):
    runs = paragraph.runs
    full_text = "".join(run.text for run in runs)
    start = full_text.find(old)
    if start == -1:
        return False
    end = start + len(old)

    pos = 0
    replaced = False
    for run in runs:
        run_start, run_end = pos, pos + len(run.text)
        pos = run_end

        if run_end <= start or run_start >= end:
            continue

        prefix = run.text[: max(0, start - run_start)]
        suffix = run.text[max(0, end - run_start):]

        if not replaced:
            run.text = prefix + new + suffix
            replaced = True
        else:
            run.text = prefix + suffix

    return True


def replace_in_row(row, old, new):
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if replace_in_paragraph(paragraph, old, new):
                return True
    return False


def fill_party_info(document, seller, buyer):
    for paragraph in document.paragraphs:
        if "BÊN A: BÊN BÁN:" in paragraph.text:
            replace_in_paragraph(
                paragraph,
                "CÔNG TY CỔ PHẦN XUẤT NHẬP KHẨU THƯƠNG MẠI ABC",
                seller["name"],
            )
        if "BÊN B: BÊN MUA" in paragraph.text:
            replace_in_paragraph(
                paragraph, "CÔNG TY TNHH TÓC XINH MỸ LỆ", buyer["name"]
            )

    table_a = document.tables[0]
    replace_in_row(table_a.rows[0], "ABCC", seller.get("address", ""))
    replace_in_row(table_a.rows[1], "0433859844", seller.get("tax_code", ""))
    replace_in_row(table_a.rows[2], "898896886", seller.get("bank_account", ""))
    replace_in_row(
        table_a.rows[3], "Ngân hàng TMCP QUÂN ĐỘI - MBANK", seller.get("bank_name", "")
    )
    replace_in_row(table_a.rows[4], "Nguyễn Thị B", seller.get("representative", ""))
    replace_in_row(table_a.rows[5], "Giám đốc", seller.get("position", ""))

    table_b = document.tables[1]
    replace_in_row(
        table_b.rows[0],
        "05 đường M2, Dự án Khu dân cư và Công viên Phước Thiện, Khu phố 28, Phường Long Bình, TP Hồ Chí Minh",
        buyer.get("address", ""),
    )
    replace_in_row(table_b.rows[1], "0319437919", buyer.get("tax_code", ""))
    replace_in_row(table_b.rows[2], "0", buyer.get("phone", ""))
    replace_in_row(table_b.rows[3], "113003051756", buyer.get("bank_account", ""))
    replace_in_row(
        table_b.rows[4],
        "Ngân hàng Thương mại Cổ phần Công Thương Việt Nam – Đồng Nai",
        buyer.get("bank_name", ""),
    )
    replace_in_row(table_b.rows[5], "Huỳnh Tấn Hải", buyer.get("representative", ""))
    replace_in_row(table_b.rows[6], "Giám đốc", buyer.get("position", ""))


def fill_contract_meta(document, contract_no, day, month, year):
    for paragraph in document.paragraphs:
        if "Số :" in paragraph.text or "Số:" in paragraph.text:
            if replace_in_paragraph(paragraph, "02062026/PPR/HĐMBHH", contract_no):
                continue
        if "Hôm nay, ngày" in paragraph.text:
            replace_in_paragraph(paragraph, "13", day)
            replace_in_paragraph(paragraph, "06", month)
            replace_in_paragraph(paragraph, "2026", year)


def _format_vnd(amount):
    return f"{amount:,}".replace(",", ".")


def _set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _set_item_row(row, item):
    cells = row.cells
    _set_cell_text(cells[0], str(item["stt"]))
    _set_cell_text(cells[1], item["description"])
    _set_cell_text(cells[2], item["unit"])
    _set_cell_text(cells[3], _format_vnd(item["quantity"]))
    _set_cell_text(cells[4], _format_vnd(item["unit_price"]))
    _set_cell_text(cells[5], _format_vnd(item["amount"]))


def fill_product_table_and_totals(document, items, totals):
    table = document.tables[2]
    rows = table.rows
    template_row = rows[1]
    subtotal_row, vat_row, total_row = rows[2], rows[3], rows[4]
    anchor_tr = subtotal_row._tr

    _set_item_row(template_row, items[0])
    for item in items[1:]:
        new_tr = deepcopy(template_row._tr)
        anchor_tr.addprevious(new_tr)
        _set_item_row(_Row(new_tr, table), item)

    _set_cell_text(subtotal_row.cells[-1], _format_vnd(totals["subtotal"]))
    for cell in vat_row.cells:
        if "VAT" in cell.text:
            replace_in_paragraph(cell.paragraphs[0], "8", totals["vat_rate"])
            break
    _set_cell_text(vat_row.cells[-1], _format_vnd(totals["vat_amount"]))
    _set_cell_text(total_row.cells[-1], _format_vnd(totals["total_payment"]))

    for paragraph in document.paragraphs:
        if "viết bằng chữ" in paragraph.text:
            replace_in_paragraph(
                paragraph,
                "Sáu mươi triệu không trăm sáu mươi chín nghìn sáu trăm đồng",
                totals["amount_in_words"],
            )


def _blacken_paragraph(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _blacken_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _blacken_paragraph(paragraph)
            for nested_table in cell.tables:
                _blacken_table(nested_table)


def force_black_text(document):
    """Set every run's font color to black, overriding any color inherited
    from the template so all generated output renders as plain black text."""
    for paragraph in document.paragraphs:
        _blacken_paragraph(paragraph)
    for table in document.tables:
        _blacken_table(table)

    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _blacken_paragraph(paragraph)
            for table in part.tables:
                _blacken_table(table)
