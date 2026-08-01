import docx
from contract_filler.principle_docx_filler import (
    fill_principle_party_a,
    fill_principle_party_b,
    fill_principle_contract_meta,
)

TEMPLATE = "tests/fixtures/principle_template.docx"

ENTITY = {
    "name": "HỘ KINH DOANH NGUYỄN VĂN TEST",
    "tax_code": "0000000001",
    "address": "123 Đường Test, Hà Nội",
    "bank_account": "",
    "representative": "Nguyễn Văn Test",
    "position": "Chủ hộ",
}

SELLER = {
    "name": "CÔNG TY CỔ PHẦN THƯƠNG MẠI XUẤT NHẬP KHẨU TKS GROUP",
    "address": "Số 9 đường Lê Văn Huấn, cụm công nghiệp Cầu Nổi, Xã Sơn Đồng, Thành phố Hà Nội, Việt Nam",
    "tax_code": "0110534607",
    "bank_account": "6062.666.88888",
    "bank_name": "Ngân hàng TMCP Quân đội – PGD Đông Đô",
    "representative": "Nguyễn Xuân Khoa",
    "position": "Giám đốc",
}


def test_fill_principle_party_a_sets_fields_and_clears_placeholder():
    document = docx.Document(TEMPLATE)

    fill_principle_party_a(document, ENTITY)

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert ENTITY["name"] in full_text
    assert "CÔNG TY TNHH TƯ VẤN PHÁT TRIỂN THỊ TRƯỜNG M.S.V" not in full_text

    table_text = "\n".join(
        "\n".join(cell.text for cell in row.cells)
        for row in document.tables[1].rows
    )
    assert ENTITY["address"] in table_text
    assert ENTITY["tax_code"] in table_text
    assert ENTITY["representative"] in table_text
    assert ENTITY["position"] in table_text
    assert "686345848" not in table_text
    assert "Ngân hàng TMCP Quân đội" not in table_text


def test_fill_principle_party_a_leaves_missing_fields_blank():
    document = docx.Document(TEMPLATE)
    entity = dict(ENTITY)
    entity["address"] = ""
    entity["position"] = ""

    fill_principle_party_a(document, entity)

    table_text = "\n".join(
        "\n".join(cell.text for cell in row.cells)
        for row in document.tables[1].rows
    )
    assert "ABC, Hà Nội" not in table_text
    assert "Giám đốc" not in table_text


def test_fill_principle_party_b_sets_fields_and_clears_placeholder():
    document = docx.Document(TEMPLATE)

    fill_principle_party_b(document, SELLER)

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert SELLER["name"] in full_text
    assert "CÔNG TY CỔ PHẦN THƯƠNG MẠI ABCXYZ" not in full_text

    table_text = "\n".join(
        "\n".join(cell.text for cell in row.cells)
        for row in document.tables[0].rows
    )
    assert SELLER["address"] in table_text
    assert SELLER["tax_code"] in table_text
    assert SELLER["representative"] in table_text
    assert SELLER["position"] in table_text
    assert "PHẠM THẾ PHONG" not in table_text
    assert "0302629806" not in table_text


def test_fill_principle_contract_meta_sets_number_and_date():
    document = docx.Document(TEMPLATE)

    fill_principle_contract_meta(document, "01/082026/HĐNT", "15", "08", "2026")

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "01/082026/HĐNT" in full_text
    assert "15/08/2026" in full_text
