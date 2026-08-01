import docx
from contract_filler.docx_filler import replace_in_paragraph


def _paragraph_with_runs(texts):
    document = docx.Document()
    paragraph = document.add_paragraph()
    for text in texts:
        paragraph.add_run(text)
    return paragraph


def test_replace_within_single_run():
    paragraph = _paragraph_with_runs(["Địa chỉ: ", "ABCC", " "])
    assert replace_in_paragraph(paragraph, "ABCC", "Số 9 đường X")
    assert paragraph.text == "Địa chỉ: Số 9 đường X "


def test_replace_spanning_multiple_runs():
    paragraph = _paragraph_with_runs(["Mã số thuế: ", "0433859", "844"])
    assert replace_in_paragraph(paragraph, "0433859844", "0110534607")
    assert paragraph.text == "Mã số thuế: 0110534607"


def test_replace_returns_false_when_not_found():
    paragraph = _paragraph_with_runs(["Hello"])
    assert replace_in_paragraph(paragraph, "Goodbye", "Hi") is False
    assert paragraph.text == "Hello"


from contract_filler.docx_filler import fill_party_info

SELLER = {
    "name": "CÔNG TY CỔ PHẦN THƯƠNG MẠI XUẤT NHẬP KHẨU TKS GROUP",
    "tax_code": "0110534607",
    "address": "Số 9 đường Lê Văn Huấn, Cụm công nghiệp Cầu Nổi, Xã Sơn Đồng, Thành phố Hà Nội, Việt Nam",
    "bank_account": "6062.666.88888",
    "bank_name": "Ngân hàng TMCP Quân đội – PGD Đông Đô",
    "representative": "Nguyễn Xuân Khoa",
    "position": "Giám đốc",
}
BUYER = {
    "name": "HỘ KINH DOANH TIỆM 81",
    "tax_code": "064200012728",
    "address": "201/65/9 Nguyễn Xí, Phường Bình Thạnh, Thành phố Hồ Chí Minh, Việt Nam",
    "bank_account": "VCB 0123456789",
    "representative": "Trần Văn Tiệm",
    "position": "Chủ hộ",
}


def test_fill_party_info_fills_seller_and_buyer_from_their_dicts():
    document = docx.Document("tests/fixtures/sample_template.docx")

    fill_party_info(document, SELLER, BUYER)

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert SELLER["name"] in full_text
    assert BUYER["name"] in full_text
    assert "ABCC" not in full_text  # old placeholder gone

    def _row_text(row):
        return "\n".join(cell.text for cell in row.cells)

    table_a_text = "\n".join(_row_text(row) for row in document.tables[0].rows)
    assert SELLER["address"] in table_a_text
    assert SELLER["tax_code"] in table_a_text
    assert SELLER["bank_account"] in table_a_text
    assert SELLER["bank_name"] in table_a_text
    assert SELLER["representative"] in table_a_text
    assert SELLER["position"] in table_a_text

    table_b_text = "\n".join(_row_text(row) for row in document.tables[1].rows)
    assert BUYER["address"] in table_b_text
    assert BUYER["tax_code"] in table_b_text
    assert BUYER["bank_account"] in table_b_text
    assert BUYER["representative"] in table_b_text
    assert BUYER["position"] in table_b_text
    assert "Ngân hàng Thương mại Cổ phần Công Thương Việt Nam – Đồng Nai" not in table_b_text  # buyer bank name always blank


def test_fill_party_info_leaves_missing_fields_blank():
    document = docx.Document("tests/fixtures/sample_template.docx")
    minimal_seller = {"name": "CÔNG TY TỐI GIẢN"}
    minimal_buyer = {"name": "HỘ KINH DOANH TỐI GIẢN"}

    fill_party_info(document, minimal_seller, minimal_buyer)

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "ABCC" not in full_text
    assert "Nguyễn Thị B" not in full_text
    assert "Huỳnh Tấn Hải" not in full_text
    assert "898896886" not in full_text
    assert "113003051756" not in full_text


from contract_filler.docx_filler import fill_contract_meta


def test_fill_contract_meta_sets_number_and_date():
    document = docx.Document("tests/fixtures/sample_template.docx")

    fill_contract_meta(document, "0099/2026/PPR/HĐMBHH", "31", "07", "2026")

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "0099/2026/PPR/HĐMBHH" in full_text
    assert "ngày 31 tháng 07 năm 2026" in full_text


from contract_filler.docx_filler import fill_product_table_and_totals

ITEMS = [
    {"stt": 1, "description": "Móc treo đồ gắn tường", "unit": "Cái",
     "quantity": 7000, "unit_price": 385, "amount": 2695000},
    {"stt": 2, "description": "Khay đựng bàn chải", "unit": "Cái",
     "quantity": 360, "unit_price": 2500, "amount": 900000},
    {"stt": 3, "description": "Bộ bàn chải đánh răng", "unit": "Bộ",
     "quantity": 300, "unit_price": 14800, "amount": 4440000},
    {"stt": 4, "description": "Dụng cụ tách sò", "unit": "Cái",
     "quantity": 100, "unit_price": 12800, "amount": 1280000},
]
TOTALS = {
    "subtotal": 9315000,
    "vat_rate": "8",
    "vat_amount": 745200,
    "total_payment": 10060200,
    "amount_in_words": "Mười triệu không trăm sáu mươi nghìn hai trăm đồng",
}


def test_fill_product_table_and_totals_writes_all_rows():
    document = docx.Document("tests/fixtures/sample_template.docx")

    fill_product_table_and_totals(document, ITEMS, TOTALS)

    table = document.tables[2]
    # header + 4 item rows + 3 totals rows
    assert len(table.rows) == 8

    for row, item in zip(table.rows[1:5], ITEMS):
        cells_text = [c.text for c in row.cells]
        assert str(item["stt"]) in cells_text[0]
        assert item["description"] in cells_text[1]
        assert item["unit"] in cells_text[2]

    assert "9.315.000" in table.rows[5].cells[-1].text
    assert "TIỀN THUẾ VAT 8%" in table.rows[6].cells[1].text
    assert "745.200" in table.rows[6].cells[-1].text
    assert "10.060.200" in table.rows[7].cells[-1].text

    full_text = "\n".join(p.text for p in document.paragraphs)
    assert TOTALS["amount_in_words"] in full_text
