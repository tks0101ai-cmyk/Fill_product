import os
import docx
from contract_filler.generator import generate_sales_contracts

PDF_FIXTURE = "tests/fixtures/sample_invoice.pdf"
TEMPLATE_FIXTURE = "tests/fixtures/sample_template.docx"
LIST_FIXTURE = "tests/fixtures/entity_list.docx"

SELLER = {
    "name": "CÔNG TY CỔ PHẦN THƯƠNG MẠI XUẤT NHẬP KHẨU TKS GROUP",
    "address": "Số 9 đường Lê Văn Huấn, cụm công nghiệp Cầu Nổi, Xã Sơn Đồng, Thành phố Hà Nội, Việt Nam",
    "tax_code": "0110534607",
    "bank_account": "6062.666.88888",
    "bank_name": "Ngân hàng TMCP Quân đội – PGD Đông Đô",
    "representative": "Nguyễn Xuân Khoa",
    "position": "Giám đốc",
}


def test_generate_sales_contracts_creates_one_file_per_entity(tmp_path):
    output_dir = str(tmp_path / "out")

    output_paths = generate_sales_contracts(
        pdf_path=PDF_FIXTURE,
        template_path=TEMPLATE_FIXTURE,
        list_path=LIST_FIXTURE,
        seller=SELLER,
        output_dir=output_dir,
        contract_no="0099/2026/PPR/HĐMBHH",
        day="31",
        month="07",
        year="2026",
    )

    assert len(output_paths) == 25
    for path in output_paths:
        assert os.path.exists(path)
        assert os.path.basename(path).startswith("HĐMB ")

    first_doc = docx.Document(output_paths[0])
    full_text = "\n".join(p.text for p in first_doc.paragraphs)

    assert SELLER["name"] in full_text
    assert "HỘ KINH DOANH TRẦN VĂN HÙNG" in full_text
    assert "0099/2026/PPR/HĐMBHH" in full_text
    assert "Mười triệu không trăm sáu mươi nghìn hai trăm đồng" in full_text

    table = first_doc.tables[2]
    assert len(table.rows) == 8
    assert "10.060.200" in table.rows[7].cells[-1].text


def test_generate_sales_contracts_fills_seller_info_in_every_output(tmp_path):
    output_dir = str(tmp_path / "out")

    output_paths = generate_sales_contracts(
        pdf_path=PDF_FIXTURE,
        template_path=TEMPLATE_FIXTURE,
        list_path=LIST_FIXTURE,
        seller=SELLER,
        output_dir=output_dir,
        contract_no="0099/2026/PPR/HĐMBHH",
        day="31",
        month="07",
        year="2026",
    )

    last_doc = docx.Document(output_paths[-1])
    table_a_text = "\n".join(
        cell.text for row in last_doc.tables[0].rows for cell in row.cells
    )
    assert SELLER["tax_code"] in table_a_text
    assert SELLER["representative"] in table_a_text
