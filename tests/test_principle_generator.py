import os
import docx
from contract_filler.principle_generator import generate_principle_contracts

LIST_FIXTURE = "tests/fixtures/entity_list.docx"
TEMPLATE_FIXTURE = "tests/fixtures/principle_template.docx"

SELLER = {
    "name": "CÔNG TY CỔ PHẦN THƯƠNG MẠI XUẤT NHẬP KHẨU TKS GROUP",
    "address": "Số 9 đường Lê Văn Huấn, cụm công nghiệp Cầu Nổi, Xã Sơn Đồng, Thành phố Hà Nội, Việt Nam",
    "tax_code": "0110534607",
    "bank_account": "6062.666.88888",
    "bank_name": "Ngân hàng TMCP Quân đội – PGD Đông Đô",
    "representative": "Nguyễn Xuân Khoa",
    "position": "Giám đốc",
}


def test_generate_principle_contracts_creates_one_file_per_entity(tmp_path):
    output_dir = str(tmp_path / "out")

    output_paths = generate_principle_contracts(
        list_path=LIST_FIXTURE,
        template_path=TEMPLATE_FIXTURE,
        output_dir=output_dir,
        seller=SELLER,
        contract_no="01/082026/HĐNT",
        day="15",
        month="08",
        year="2026",
    )

    assert len(output_paths) == 25
    for path in output_paths:
        assert os.path.exists(path)
        assert os.path.basename(path).startswith("HĐNT ")

    first_doc = docx.Document(output_paths[0])
    full_text = "\n".join(p.text for p in first_doc.paragraphs)
    assert "HỘ KINH DOANH TRẦN VĂN HÙNG" in full_text
    assert SELLER["name"] in full_text
    assert "01/082026/HĐNT" in full_text


def test_generate_principle_contracts_leaves_missing_fields_blank(tmp_path):
    output_dir = str(tmp_path / "out")

    output_paths = generate_principle_contracts(
        list_path=LIST_FIXTURE,
        template_path=TEMPLATE_FIXTURE,
        output_dir=output_dir,
        seller=SELLER,
        contract_no="01/082026/HĐNT",
        day="15",
        month="08",
        year="2026",
    )

    first_doc = docx.Document(output_paths[0])
    table_text = "\n".join(
        cell.text for row in first_doc.tables[1].rows for cell in row.cells
    )
    assert "686345848" not in table_text  # HKD entity has no bank_account
